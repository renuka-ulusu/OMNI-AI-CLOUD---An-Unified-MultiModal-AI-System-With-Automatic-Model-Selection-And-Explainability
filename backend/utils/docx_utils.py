from pathlib import Path
from typing import Dict
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document

from backend.utils.pdf_utils import normalize_utf8_text


MAX_DOCX_SIZE_BYTES = 25 * 1024 * 1024
MAX_DOCX_PARAGRAPHS = 10000


def validate_docx_file(docx_path: str) -> Dict[str, int]:
    """
    Strict DOCX validation:
    - extension must be .docx
    - file exists and non-empty
    - max file size
    - valid ZIP container with required DOCX internals
    - parseable by python-docx
    - paragraph count within limit
    """
    path = Path(docx_path)

    if path.suffix.lower() != ".docx":
        raise ValueError("Only .docx Word documents are supported")

    if not path.exists() or not path.is_file():
        raise ValueError("Uploaded DOCX file not found")

    file_size = path.stat().st_size
    if file_size <= 0:
        raise ValueError("Uploaded DOCX is empty")

    if file_size > MAX_DOCX_SIZE_BYTES:
        raise ValueError(
            f"DOCX exceeds maximum allowed size of {MAX_DOCX_SIZE_BYTES // (1024 * 1024)}MB"
        )

    with open(path, "rb") as file_obj:
        signature = file_obj.read(2)

    if signature != b"PK":
        raise ValueError("Invalid DOCX signature. Upload a valid Word .docx file")

    if not is_zipfile(path):
        raise ValueError("Invalid DOCX container")

    try:
        with ZipFile(path, "r") as archive:
            names = set(archive.namelist())
    except BadZipFile as error:
        raise ValueError(f"Invalid DOCX archive: {error}") from error

    required_parts = {"[Content_Types].xml", "word/document.xml"}
    if not required_parts.issubset(names):
        raise ValueError("Invalid DOCX structure")

    try:
        document = Document(str(path))
    except Exception as error:
        raise ValueError(f"Unreadable DOCX document: {error}") from error

    paragraph_count = len(document.paragraphs)
    if paragraph_count > MAX_DOCX_PARAGRAPHS:
        raise ValueError(f"DOCX exceeds maximum paragraph limit of {MAX_DOCX_PARAGRAPHS}")

    return {
        "file_size": file_size,
        "paragraph_count": paragraph_count,
    }


def extract_text_from_docx(docx_path: str) -> str:
    """
    Extract text from DOCX paragraphs and table cells.
    Enforces strict UTF-8 normalization.
    """
    document = Document(docx_path)

    text_blocks = []

    for paragraph in document.paragraphs:
        value = paragraph.text.strip()
        if value:
            text_blocks.append(value)

    for table in document.tables:
        for row in table.rows:
            cell_values = []
            for cell in row.cells:
                cell_text = " ".join(
                    p.text.strip() for p in cell.paragraphs if p.text and p.text.strip()
                )
                if cell_text:
                    cell_values.append(cell_text)
            if cell_values:
                text_blocks.append(" | ".join(cell_values))

    extracted = "\n".join(text_blocks).strip()
    return normalize_utf8_text(extracted)
